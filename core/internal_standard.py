import docx
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import re
import sys
import os

# Sửa lỗi hiển thị tiếng Việt trên Terminal Windows
if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

# Từ điển sửa lỗi chính tả và các cụm từ nội bộ của TEXO
TEXO_SPELL_CHECK_MAP = {
    r"\bquy ttrình\b": "quy trình",
    r"\bloà việc\b": "làm việc",
    r"\bsu phat trien\b": "sự phát triển",
    r"\bsai xót\b": "sai sót",
    r"\bko\b": "không",
    r"\bchình ký\b": "trình ký",
    r"\bthe thức\b": "thể thức",
    r"\btu vấn\b": "tư vấn",
    r"\bdu án\b": "dự án",
    # Quy tắc ngữ pháp (Dấu phẩy sau trạng/liên từ)
    r"\b(Tuy nhiên|Do đó|Bởi vậy|Vì vậy|Đến nay|Hiện nay|Theo đó)\s+([^,])": r"\1, \2",
    # Quy tắc chính tả (l/n, tr/ch...) - Ví dụ một số cụm hay sai
    r"\blên kế hoạch\b": "lên kế hoạch",
    r"\bnên kế hoạch\b": "lên kế hoạch",
}

def correct_grammar_and_spell(text):
    # 1. Áp dụng từ điển chính tả & ngữ pháp (dấu phẩy sau Tuy nhiên...)
    for pattern, replacement in TEXO_SPELL_CHECK_MAP.items():
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
    
    # 2. Quy tắc viết ngày tháng (Chỉ áp dụng cho Ngày/Tháng để tránh sai lệch các số thứ tự khác)
    # Định dạng: ngày 1 -> ngày 01, tháng 2 -> tháng 02
    text = re.sub(r"(ngày|tháng)\s+([1-9])\b", r"\1 0\2", text, flags=re.IGNORECASE)
    
    # Định dạng: 1/2/2026 -> 01/02/2026
    def date_slash_fixer(match):
        d, m, y = match.groups()
        new_d = d if len(d) == 2 else "0" + d
        new_m = m if len(m) == 2 else "0" + m
        return f"{new_d}/{new_m}/{y}"
    
    text = re.sub(r"\b(\d{1,2})/(\d{1,2})/(\d{4})\b", date_slash_fixer, text)

    # 3. Quy tắc "bốn" vs "tư" (Đơn giản hóa: nếu trước nó là "mươi" thì dùng "tư", trừ mười bốn)
    text = re.sub(r"mươi bốn", "mươi tư", text, flags=re.IGNORECASE)
    
    return text

def apply_texo_internal_standard(input_path, output_path, is_letterhead=False):
    doc = docx.Document(input_path)
    
    # 1. Căn lề khổ giấy A4 & Margins
    for section in doc.sections:
        # A4 size: 21.0cm x 29.7cm
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        
        if is_letterhead:
            # Letterhead (Logo chìm): T40, B30, L30, R20
            section.top_margin = Mm(40)
            section.bottom_margin = Mm(30)
            section.left_margin = Mm(30)
            section.right_margin = Mm(20)
        else:
            # Giấy thường: T20, B20, L30, R20
            section.top_margin = Mm(20)
            section.bottom_margin = Mm(20)
            section.left_margin = Mm(30)
            section.right_margin = Mm(20)
    
    # 2. Xử lý các đoạn văn bản (Paragraphs)
    for para in doc.paragraphs:
        # Căn lề đều hai bên (trừ Center)
        if para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # --- CẤU HÌNH SPACING (6pt/3pt) ---
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(3)
        # Loại bỏ "Don't add space between paragraphs of same style" (nếu có hỗ trợ)
        try:
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            para.paragraph_format.line_spacing = Pt(17)
            # Tắt tính năng contextual spacing
            pPr = para._element.get_or_add_pPr()
            contextualSpacer = pPr.find(docx.oxml.ns.qn('w:contextualSpacing'))
            if contextualSpacer is not None:
                pPr.remove(contextualSpacer)
        except:
            pass

        # Kiểm tra xem có phải "Dòng tiêu đề lớn" (TỜ TRÌNH, CÔNG VĂN...) không
        is_big_title = False
        clean_text = para.text.strip()
        if clean_text.isupper() and len(clean_text) < 100:
            is_big_title = True

        # Xử lý nội dung và format ở cấp độ RUN
        for run in para.runs:
            if run.text.strip():
                run.text = correct_grammar_and_spell(run.text)
                
            # Áp dụng Font chuẩn
            run.font.name = 'Times New Roman'
            
            # --- CẤU HÌNH SIZE THEO YÊU CẦU ANH VŨ ---
            if is_big_title:
                run.font.size = Pt(14)  # Tiêu đề lớn dùng 14
            elif run.text.isupper():
                run.font.size = Pt(12)  # Các tiêu đề con/khác in hoa dùng 12
            else:
                run.font.size = Pt(13)  # Nội dung thường dùng 13
            
            # Giữ nguyên Bold/Italic, auto-bold "Kính gửi"
            if "Kính gửi" in run.text:
                run.bold = True

    # 3. Xử lý bảng biểu
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Phải lặp qua cả các đoạn văn trong cell
                for para in cell.paragraphs:
                    # Cấu hình Spacing và Indentation để bảng nhìn thoáng hơn
                    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    para.paragraph_format.line_spacing = Pt(15) 
                    para.paragraph_format.space_before = Pt(6)
                    para.paragraph_format.space_after = Pt(3)
                    
                    # Thêm 2mm lề để chữ không bị dính sát vào cạnh bảng
                    para.paragraph_format.left_indent = Mm(2)
                    para.paragraph_format.right_indent = Mm(2)
                    
                    for run in para.runs:
                        if run.text.strip():
                            # Sửa lỗi chính tả & Ngày/Tháng trong bảng
                            run.text = correct_grammar_and_spell(run.text)
                        
                        # Áp dụng font chuẩn cho bảng
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12) 
    
    # 4. Lưu file
    doc.save(output_path)
    print(f"--- Đã áp dụng CHUẨN NỘI BỘ TEXO và lưu tại: {output_path} ---")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        input_file = sys.argv[1]
        letterhead_mode = False
        
        if "--letterhead" in sys.argv:
            letterhead_mode = True
        elif "--normal" in sys.argv:
            letterhead_mode = False
        else:
            # Hỏi người dùng trực tiếp qua console
            print("Bạn sử dụng loại giấy nào?")
            print("1. Giấy Letterhead (có logo chìm - Lề 40/30/30/20)")
            print("2. Giấy thường (không logo - Lề 20/20/30/20)")
            choice = input("Lựa chọn (1 hoặc 2): ").strip()
            if choice == "1":
                letterhead_mode = True
            else:
                letterhead_mode = False

        dir_name = os.path.dirname(input_file)
        if not dir_name: dir_name = "."
        base_name = os.path.basename(input_file)
        output_file = os.path.join(dir_name, "TEXO_Standardized_" + base_name)
        
        apply_texo_internal_standard(input_file, output_file, is_letterhead=letterhead_mode)
    else:
        print("Sử dụng: python texo_internal_standard.py <path> [--letterhead|--normal]")
