import docx
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import sys
import os

# Sửa lỗi hiển thị tiếng Việt trên Terminal Windows
if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

# Từ điển sửa lỗi chính tả và viết tắt phổ biến của TEXO
VIETNAMESE_SPELL_CHECK_MAP = {
    r"\bquy ttrình\b": "quy trình",
    r"\bloà việc\b": "làm việc",
    r"\bsu phat trien\b": "sự phát triển",
    r"\bsai xót\b": "sai sót",
    r"\bko\b": "không",
    r"\bchình ký\b": "trình ký",
    r"\bthe thức\b": "thể thức",
    r"\btu vấn\b": "tư vấn",
    r"\bdu án\b": "dự án",
}

def correct_text(text):
    # 1. Sửa lỗi từ điển
    for pattern, replacement in VIETNAMESE_SPELL_CHECK_MAP.items():
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
    
    # 2. Sửa lỗi Ngày/Tháng thiếu số 0 (chỉ áp dụng cho Ngày/Tháng)
    text = re.sub(r"(ngày|tháng)\s+([1-9])\b", r"\1 0\2", text, flags=re.IGNORECASE)
    
    return text

def apply_nd30_standard(input_path, output_path):
    doc = docx.Document(input_path)
    
    # 1. Page Size A4 & Margins (NĐ 30/2020: T20, B20, L30, R20)
    for section in doc.sections:
        # A4 size: 21.0cm x 29.7cm
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(30)
        section.right_margin = Mm(20) # Sửa từ 15 -> 20 cho chuẩn hơn
    
    # 2. Iterate through paragraphs
    for para in doc.paragraphs:
        # Sửa chính tả
        if para.text.strip():
            for run in para.runs:
                run.text = correct_text(run.text)

        # Căn lề đều hai bên
        if para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        para.paragraph_format.line_spacing = 1.15
        
        for run in para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            
    # Xử lý các bảng biểu (Tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(13)
    
    doc.save(output_path)
    print(f"--- Đã áp dụng chuẩn NĐ 30/2020 v\u00e0 l\u01b0u t\u1ea1i: {output_path} ---")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        input_file = sys.argv[1]
        dir_name = os.path.dirname(input_file)
        base_name = os.path.basename(input_file)
        # Nếu không lấy được dir_name (file nằm ở root hiện tại), lưu vào thư mục hiện tại
        if not dir_name: dir_name = "."
        output_file = os.path.join(dir_name, "Standardized_" + base_name)
        
        apply_nd30_standard(input_file, output_file)
    else:
        print("Vui lòng cung cấp đường dẫn file .docx")
